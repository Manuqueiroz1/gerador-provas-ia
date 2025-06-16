# === GERADOR DE PROVAS IA - VERSÃO CORRIGIDA ===
# Mantém toda funcionalidade original + correções específicas

import streamlit as st
import openai
import time
import os
import re
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# === CONFIGURAÇÃO SEGURA DA API ===
if "OPENAI_API_KEY" in st.secrets:
    OPENAI_API_KEY = st.secrets["OPENAI_API_KEY"]
    ASSISTANT_ID = st.secrets["ASSISTANT_ID"]
else:
    st.error("🔑 Configure suas chaves da API nas configurações do Streamlit!")
    st.info("Acesse as configurações do app para adicionar as chaves secretas.")
    st.stop()

openai.api_key = OPENAI_API_KEY

st.set_page_config(
    page_title="Gerador de Provas IA Completo", 
    layout="wide", 
    initial_sidebar_state="expanded",
    page_icon="📚"
)

# === INICIALIZAR SESSION STATE ===
if 'prova_gerada' not in st.session_state:
    st.session_state.prova_gerada = None
if 'documentos_prontos' not in st.session_state:
    st.session_state.documentos_prontos = None
if 'imagens_geradas' not in st.session_state:
    st.session_state.imagens_geradas = []

# === LISTA EXTENSA DE TEMAS POR SÉRIE ===
TEMAS_COMPLETOS = {
    "6º Ano": [
        "Números Naturais", "Operações Básicas", "Frações", "Decimais", 
        "Porcentagem Básica", "Geometria Plana Básica", "Perímetro e Área",
        "Unidades de Medida", "Sistema Monetário", "Gráficos e Tabelas",
        "Múltiplos e Divisores", "Números Primos", "Expressões Numéricas",
        "Ângulos", "Polígonos", "Simetria"
    ],
    "7º Ano": [
        "Números Inteiros", "Números Racionais", "Equações do 1º Grau",
        "Inequações", "Razão e Proporção", "Regra de Três", "Porcentagem",
        "Geometria: Triângulos", "Quadriláteros", "Circunferência",
        "Estatística Básica", "Probabilidade", "Expressões Algébricas",
        "Plano Cartesiano", "Ângulos em Polígonos", "Teorema de Tales"
    ],
    "8º Ano": [
        "Sistemas de Equações Lineares", "Produtos Notáveis", "Fatoração",
        "Frações Algébricas", "Função do 1º Grau", "Teorema de Pitágoras",
        "Áreas e Volumes", "Semelhança de Triângulos", "Relações Métricas",
        "Dízimas Periódicas", "Potenciação", "Radiciação",
        "Monômios e Polinômios", "Geometria Analítica Básica"
    ],
    "9º Ano": [
        "Função Quadrática", "Equações do 2º Grau", "Trigonometria no Triângulo Retângulo",
        "Relações Métricas na Circunferência", "Razões Trigonométricas",
        "Geometria Espacial", "Estatística e Probabilidade", "Progressões",
        "Sistemas de Inequações", "Função Exponencial Básica", "Logaritmos Básicos",
        "Matemática Financeira", "Análise Combinatória Básica", "Semelhança de Triângulos"
    ],
    "1º Ano EM": [
        "Conjuntos", "Funções", "Função Afim", "Função Quadrática",
        "Função Exponencial", "Função Logarítmica", "Progressões Aritméticas",
        "Progressões Geométricas", "Trigonometria", "Geometria Plana",
        "Estatística", "Análise Combinatória", "Probabilidade",
        "Matemática Financeira", "Sistemas Lineares"
    ],
    "2º Ano EM": [
        "Matrizes", "Determinantes", "Sistemas Lineares", "Geometria Espacial",
        "Geometria Analítica", "Circunferência", "Elipse", "Hipérbole",
        "Parábola", "Números Complexos", "Polinômios", "Equações Polinomiais",
        "Binômio de Newton", "Probabilidade Avançada", "Estatística Avançada"
    ],
    "3º Ano EM": [
        "Geometria Espacial Avançada", "Geometria Analítica Espacial",
        "Limites", "Derivadas Básicas", "Integrais Básicas", "Funções Trigonométricas",
        "Análise Combinatória Avançada", "Probabilidade Condicional",
        "Estatística Inferencial", "Matemática Financeira Avançada",
        "Sequências e Séries", "Tópicos de Cálculo"
    ]
}

# === SISTEMA DE LIMPEZA DE FORMATAÇÃO MATEMÁTICA ===
def limpar_formatacao_latex(texto):
    """Remove ou converte códigos LaTeX para texto limpo"""
    
    # Remover delimitadores LaTeX
    texto = re.sub(r'\\\(([^)]*)\\\)', r'\1', texto)
    texto = re.sub(r'\\\[([^\]]*)\\\]', r'\1', texto)
    
    # Converter frações
    texto = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1)/(\2)', texto)
    
    # Converter raízes
    texto = re.sub(r'\\sqrt\{([^}]+)\}', r'√(\1)', texto)
    
    # Converter exponenciais
    texto = re.sub(r'\^2\b', '²', texto)
    texto = re.sub(r'\^3\b', '³', texto)
    texto = re.sub(r'\^4\b', '⁴', texto)
    texto = re.sub(r'\^{2}', '²', texto)
    texto = re.sub(r'\^{3}', '³', texto)
    texto = re.sub(r'\^{([^}]+)}', r'^(\1)', texto)
    
    # Limpar outros símbolos LaTeX comuns
    substituicoes = {
        r'\\cdot': '·',
        r'\\times': '×',
        r'\\div': '÷',
        r'\\pm': '±',
        r'\\pi': 'π',
        r'\\alpha': 'α',
        r'\\beta': 'β',
        r'\\gamma': 'γ',
        r'\\theta': 'θ',
        r'\\leq': '≤',
        r'\\geq': '≥',
        r'\\neq': '≠',
        r'\\approx': '≈'
    }
    
    for latex, simbolo in substituicoes.items():
        texto = re.sub(latex, simbolo, texto)
    
    # Limpar espaços extras
    texto = re.sub(r'\s+', ' ', texto)
    texto = texto.strip()
    
    return texto

# === SISTEMA DE ANÁLISE CONTEXTUAL INTELIGENTE ===
def analisar_contexto_questao_especifica(enunciado, numero_questao):
    """Analisa o contexto específico de cada questão individual"""
    
    contexto = {
        'precisa_imagem': False,
        'tipo_visualizacao': None,
        'dados_especificos': {},
        'nivel_complexidade': 'simples'
    }
    
    # Extrair dados específicos do enunciado
    medidas = re.findall(r'(\d+(?:\.\d+)?)\s*(?:cm|metros?|km)', enunciado.lower())
    angulos = re.findall(r'(\d+)°', enunciado)
    coordenadas = re.findall(r'\((-?\d+),\s*(-?\d+)\)', enunciado)
    razoes = re.findall(r'(\d+:\d+)', enunciado)
    
    # TRIGONOMETRIA - Análise específica
    if any(palavra in enunciado.lower() for palavra in ['seno', 'coseno', 'tangente', 'hipotenusa', 'cateto']):
        contexto.update({
            'precisa_imagem': True,
            'tipo_visualizacao': 'triangulo_trigonometrico',
            'dados_especificos': {
                'medidas': medidas,
                'angulos': angulos,
                'tipo_problema': 'trigonometria'
            },
            'nivel_complexidade': 'detalhado'
        })
    
    # SEMELHANÇA DE TRIÂNGULOS
    elif 'semelhança' in enunciado.lower() or 'semelhantes' in enunciado.lower():
        contexto.update({
            'precisa_imagem': True,
            'tipo_visualizacao': 'triangulos_semelhantes',
            'dados_especificos': {
                'medidas': medidas,
                'razoes': razoes,
                'tipo_problema': 'semelhanca'
            }
        })
    
    # GEOMETRIA ANALÍTICA
    elif coordenadas or 'plano cartesiano' in enunciado.lower() or 'coordenadas' in enunciado.lower():
        contexto.update({
            'precisa_imagem': True,
            'tipo_visualizacao': 'plano_cartesiano',
            'dados_especificos': {
                'coordenadas': coordenadas,
                'tipo_problema': 'geometria_analitica'
            }
        })
    
    # SISTEMAS LINEARES
    elif 'sistema' in enunciado.lower() and any(op in enunciado for op in ['x +', 'y =', 'x -', 'y +']):
        contexto.update({
            'precisa_imagem': True,
            'tipo_visualizacao': 'sistema_linear',
            'dados_especificos': {
                'equacoes': extrair_equacoes(enunciado),
                'tipo_problema': 'sistema_linear'
            }
        })
    
    # FUNÇÕES
    elif 'função' in enunciado.lower() or 'f(x)' in enunciado or 'y =' in enunciado:
        contexto.update({
            'precisa_imagem': True,
            'tipo_visualizacao': 'grafico_funcao',
            'dados_especificos': {
                'funcoes': extrair_funcoes(enunciado),
                'tipo_problema': 'funcao'
            }
        })
    
    return contexto

def extrair_equacoes(texto):
    """Extrai equações do texto"""
    equacoes = re.findall(r'[xy]\s*[+\-]\s*[xy]?\s*=\s*\d+', texto)
    return equacoes

def extrair_funcoes(texto):
    """Extrai funções do texto"""
    funcoes = re.findall(r'[yf]\s*\([x]\)?\s*=\s*[^,\n\.]+', texto)
    return funcoes

# === GERADORES DE IMAGEM ESPECÍFICOS ===
def gerar_triangulo_trigonometrico_especifico(questao_num, dados):
    """Gera triângulo com dados específicos da questão"""
    try:
        plt.ioff()
        fig, ax = plt.subplots(figsize=(10, 8))
        
        medidas = dados.get('medidas', [])
        angulos = dados.get('angulos', [])
        
        # Usar medidas reais da questão ou padrão
        if len(medidas) >= 2:
            base = float(medidas[0])
            altura = float(medidas[1])
        else:
            base, altura = 4, 3
        
        # Vértices do triângulo retângulo
        vertices = np.array([[0, 0], [base, 0], [base, altura]])
        triangle = plt.Polygon(vertices, fill=False, edgecolor='blue', linewidth=3)
        ax.add_patch(triangle)
        
        # Vértices
        ax.plot([0, base, base], [0, 0, altura], 'ro', markersize=8)
        ax.text(-0.2, -0.2, 'A', fontsize=12, weight='bold')
        ax.text(base+0.1, -0.2, 'B', fontsize=12, weight='bold')
        ax.text(base+0.1, altura+0.1, 'C', fontsize=12, weight='bold')
        
        # Ângulo reto
        square = plt.Rectangle((base-0.3, 0), 0.3, 0.3, fill=False, linewidth=2)
        ax.add_patch(square)
        
        # Ângulo específico da questão
        if angulos:
            angulo = float(angulos[0])
            angle_arc = patches.Arc((0, 0), 1, 1, angle=0, theta1=0, theta2=angulo, 
                                   color='red', linewidth=2)
            ax.add_patch(angle_arc)
            ax.text(0.3, 0.1, f'{angulo}°', fontsize=12, color='red', weight='bold')
        
        # Labels com medidas reais
        if medidas:
            ax.text(base/2, -0.3, f'{medidas[0]} cm', ha='center', fontsize=11, weight='bold')
            if len(medidas) > 1:
                ax.text(base+0.3, altura/2, f'{medidas[1]} cm', ha='center', fontsize=11, weight='bold')
        
        # Razões trigonométricas
        ax.text(base/2, -0.6, 'cateto adjacente', ha='center', fontsize=10, style='italic')
        ax.text(base+0.8, altura/2, 'cateto oposto', ha='center', fontsize=10, rotation=90, style='italic')
        ax.text(base/2-0.5, altura/2+0.3, 'hipotenusa', ha='center', fontsize=10, rotation=37, style='italic')
        
        ax.set_xlim(-1, base+2)
        ax.set_ylim(-1, altura+2)
        ax.set_aspect('equal')
        ax.axis('off')
        ax.set_title(f'Questão {questao_num}: Trigonometria', fontsize=14, weight='bold')
        
        filename = f"questao_{questao_num}_trigonometria.png"
        plt.tight_layout()
        plt.savefig(filename, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        
        return filename
        
    except Exception as e:
        print(f"Erro ao gerar imagem trigonométrica: {e}")
        return None

def gerar_triangulos_semelhantes_especificos(questao_num, dados):
    """Gera dois triângulos semelhantes com dados da questão"""
    try:
        plt.ioff()
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 7))
        
        medidas = dados.get('medidas', [])
        razoes = dados.get('razoes', [])
        
        # Triângulo 1 - menor
        if len(medidas) >= 2:
            lado1, lado2 = float(medidas[0]), float(medidas[1])
        else:
            lado1, lado2 = 3, 4
        
        vertices1 = np.array([[0, 0], [lado1, 0], [lado1*0.8, lado2]])
        triangle1 = plt.Polygon(vertices1, fill=False, edgecolor='blue', linewidth=3)
        ax1.add_patch(triangle1)
        
        # Labels
        ax1.text(lado1/2, -0.3, f'{lado1} cm', ha='center', fontsize=11, weight='bold')
        ax1.text(lado1+0.3, lado2/2, f'{lado2} cm', ha='center', fontsize=11, weight='bold')
        ax1.set_title('Triângulo ABC', fontsize=12, weight='bold')
        ax1.set_aspect('equal')
        ax1.axis('off')
        
        # Triângulo 2 - maior (proporção baseada na questão)
        if razoes:
            razao_str = razoes[0]
            if ':' in razao_str:
                r1, r2 = razao_str.split(':')
                fator = float(r2) / float(r1)
            else:
                fator = 2
        else:
            fator = 2
        
        lado1_2, lado2_2 = lado1 * fator, lado2 * fator
        vertices2 = np.array([[0, 0], [lado1_2, 0], [lado1_2*0.8, lado2_2]])
        triangle2 = plt.Polygon(vertices2, fill=False, edgecolor='red', linewidth=3)
        ax2.add_patch(triangle2)
        
        ax2.text(lado1_2/2, -0.3, f'{lado1_2} cm', ha='center', fontsize=11, weight='bold')
        ax2.text(lado1_2+0.3, lado2_2/2, f'{lado2_2} cm', ha='center', fontsize=11, weight='bold')
        ax2.set_title('Triângulo DEF', fontsize=12, weight='bold')
        ax2.set_aspect('equal')
        ax2.axis('off')
        
        plt.suptitle(f'Questão {questao_num}: Triângulos Semelhantes', fontsize=14, weight='bold')
        
        filename = f"questao_{questao_num}_semelhantes.png"
        plt.tight_layout()
        plt.savefig(filename, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        
        return filename
        
    except Exception as e:
        print(f"Erro ao gerar triângulos semelhantes: {e}")
        return None

def gerar_plano_cartesiano_especifico(questao_num, dados):
    """Gera plano cartesiano com coordenadas específicas"""
    try:
        plt.ioff()
        fig, ax = plt.subplots(figsize=(10, 8))
        
        coordenadas = dados.get('coordenadas', [])
        
        if coordenadas:
            # Usar coordenadas reais da questão
            pontos_x = [int(coord[0]) for coord in coordenadas]
            pontos_y = [int(coord[1]) for coord in coordenadas]
            
            x_min, x_max = min(pontos_x) - 2, max(pontos_x) + 2
            y_min, y_max = min(pontos_y) - 2, max(pontos_y) + 2
        else:
            # Padrão se não tiver coordenadas
            pontos_x, pontos_y = [1, 3], [2, 4]
            x_min, x_max, y_min, y_max = -1, 5, -1, 5
        
        ax.set_xlim(x_min, x_max)
        ax.set_ylim(y_min, y_max)
        ax.grid(True, alpha=0.3)
        ax.axhline(y=0, color='k', linewidth=1)
        ax.axvline(x=0, color='k', linewidth=1)
        
        # Plotar pontos específicos
        for i, (x, y) in enumerate(zip(pontos_x, pontos_y)):
            ax.plot(x, y, 'ro', markersize=10)
            ax.annotate(f'({x},{y})', (x, y), xytext=(x+0.2, y+0.2), fontsize=12, weight='bold')
        
        # Se houver 2 pontos, desenhar segmento
        if len(pontos_x) == 2:
            ax.plot(pontos_x, pontos_y, 'b-', linewidth=2)
        
        ax.set_xlabel('x', fontsize=12)
        ax.set_ylabel('y', fontsize=12)
        ax.set_title(f'Questão {questao_num}: Plano Cartesiano', fontsize=14, weight='bold')
        
        filename = f"questao_{questao_num}_cartesiano.png"
        plt.tight_layout()
        plt.savefig(filename, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        
        return filename
        
    except Exception as e:
        print(f"Erro ao gerar plano cartesiano: {e}")
        return None

def gerar_imagem_contextual_especifica(questao_num, enunciado, tema):
    """Sistema principal que decide qual tipo de imagem gerar baseado no contexto específico"""
    
    # Analisar contexto específico da questão
    contexto = analisar_contexto_questao_especifica(enunciado, questao_num)
    
    if not contexto['precisa_imagem']:
        return None
    
    # Roteamento para gerador específico
    try:
        if contexto['tipo_visualizacao'] == 'triangulo_trigonometrico':
            return gerar_triangulo_trigonometrico_especifico(questao_num, contexto['dados_especificos'])
        
        elif contexto['tipo_visualizacao'] == 'triangulos_semelhantes':
            return gerar_triangulos_semelhantes_especificos(questao_num, contexto['dados_especificos'])
        
        elif contexto['tipo_visualizacao'] == 'plano_cartesiano':
            return gerar_plano_cartesiano_especifico(questao_num, contexto['dados_especificos'])
        
        else:
            return None
            
    except Exception as e:
        st.warning(f"Erro ao gerar imagem para questão {questao_num}: {str(e)}")
        return None

# === SISTEMA DE PARSING DE QUESTÕES ===
def parse_questao_individual(questao_text):
    """Extrai dados de uma questão individual"""
    dados = {"titulo": "", "enunciado": "", "alternativas": "", "resolucao": "", "referencia": ""}
    
    # Extrair título/número da questão
    titulo_match = re.search(r"##\s*Questão\s*(\d+):?\s*(.*?)(?:\n|$)", questao_text, re.IGNORECASE)
    if titulo_match:
        dados["numero"] = titulo_match.group(1)
        dados["titulo"] = titulo_match.group(2).strip()
    
    # Extrair enunciado
    enunciado_match = re.search(r"\*\*Enunciado:\*\*\s*(.*?)(?=\*\*Alternativas:\*\*|\*\*Resolução:\*\*|\*\*Referência:\*\*|##|$)", 
                               questao_text, re.DOTALL | re.IGNORECASE)
    if enunciado_match:
        dados["enunciado"] = enunciado_match.group(1).strip()
    
    # Extrair alternativas
    alt_match = re.search(r"\*\*Alternativas:\*\*\s*(.*?)(?=\*\*Resolução:\*\*|\*\*Referência:\*\*|##|$)", 
                         questao_text, re.DOTALL | re.IGNORECASE)
    if alt_match:
        dados["alternativas"] = alt_match.group(1).strip()
    
    # Extrair resolução
    resolucao_match = re.search(r"\*\*Resolução:\*\*\s*(.*?)(?=\*\*Referência:\*\*|##|$)", 
                               questao_text, re.DOTALL | re.IGNORECASE)
    if resolucao_match:
        dados["resolucao"] = resolucao_match.group(1).strip()
    
    # Extrair referência
    referencia_match = re.search(r"\*\*Referência:\*\*\s*(.*?)(?=##|$)", 
                                questao_text, re.DOTALL | re.IGNORECASE)
    if referencia_match:
        dados["referencia"] = referencia_match.group(1).strip()
    
    return dados

def parse_prova_completa(texto):
    """Extrai todas as questões e seções da prova"""
    
    # Dividir por questões
    blocos = re.split(r"\n##\s*Questão", texto, flags=re.IGNORECASE)
    
    # Primeira parte é introdução
    introducao = ""
    questoes = []
    
    if blocos:
        introducao = blocos[0].strip()
        
        # Processar cada questão
        for i, bloco in enumerate(blocos[1:], 1):
            questao_text = f"## Questão {bloco.strip()}"
            dados = parse_questao_individual(questao_text)
            
            if dados["enunciado"].strip():
                dados["numero"] = i
                questoes.append(dados)
    
    return introducao, questoes

# === CRIAÇÃO DE DOCUMENTOS ===
def criar_documento_prova_completo(texto, serie, tema, incluir_imagens=True):
    """Cria documento DOCX com estrutura completa"""
    try:
        document = Document()
        
        # Cabeçalho principal
        titulo = document.add_heading(f"PROVA DE MATEMÁTICA - {serie.upper()}", level=0)
        titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        subtitulo = document.add_heading(f"Tema: {tema}", level=1)
        subtitulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        document.add_paragraph("")
        
        # Dados do aluno
        p_dados = document.add_paragraph("Data: ___/___/_____     Nome: _________________________     Turma: _____")
        p_dados.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        document.add_paragraph("")
        
        # Processar conteúdo
        introducao, questoes = parse_prova_completa(texto)
        
        # Adicionar instruções se existirem
        if introducao:
            # Extrair instruções
            instrucoes_match = re.search(r"Instruções:.*?(?=##|$)", introducao, re.DOTALL | re.IGNORECASE)
            if instrucoes_match:
                p_inst = document.add_paragraph(instrucoes_match.group(0))
                p_inst.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                document.add_paragraph("")
        
        if not questoes:
            st.warning("⚠️ Nenhuma questão válida encontrada.")
            return None
        
        # Limpar imagens antigas
        for img in st.session_state.imagens_geradas:
            if os.path.exists(img):
                try:
                    os.remove(img)
                except:
                    pass
        st.session_state.imagens_geradas = []
        
        # QUESTÕES COM IMAGENS CONTEXTUAIS
        for idx, questao in enumerate(questoes, start=1):
            # Título da questão
            p_titulo = document.add_paragraph()
            p_titulo.add_run(f"{idx}) ").bold = True
            p_titulo.add_run(questao["enunciado"])
            
            # GERAR IMAGEM CONTEXTUAL ESPECÍFICA
            if incluir_imagens and questao["enunciado"]:
                with st.spinner(f"🎨 Gerando imagem contextual para questão {idx}..."):
                    img_file = gerar_imagem_contextual_especifica(idx, questao["enunciado"], tema)
                    
                    if img_file and os.path.exists(img_file):
                        try:
                            # Adicionar imagem ao documento
                            p_img = document.add_paragraph()
                            p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            document.add_picture(img_file, width=Inches(5))
                            document.add_paragraph("")
                            
                            # Salvar na lista para limpeza posterior
                            st.session_state.imagens_geradas.append(img_file)
                            
                            st.success(f"✅ Imagem contextual gerada para questão {idx}")
                            
                        except Exception as e:
                            st.warning(f"⚠️ Erro ao adicionar imagem da questão {idx}: {str(e)}")
                    else:
                        st.info(f"ℹ️ Questão {idx}: Não necessita imagem ou erro na geração")
            
            # Alternativas
            if questao["alternativas"]:
                alternativas_linhas = questao["alternativas"].split('\n')
                for alt in alternativas_linhas:
                    if alt.strip():
                        document.add_paragraph(alt.strip())
            
            document.add_paragraph("")
        
        # Salvar documento
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        nome_arquivo = f"prova_{serie.replace(' ', '_')}_{tema.replace(' ', '_')}_{timestamp}.docx"
        document.save(nome_arquivo)
        
        return nome_arquivo
        
    except Exception as e:
        st.error(f"Erro ao criar documento: {str(e)}")
        return None

def criar_documento_gabarito(texto, serie, tema):
    """Cria documento separado com gabarito e resoluções"""
    try:
        document = Document()
        
        # Cabeçalho
        titulo = document.add_heading(f"GABARITO E RESOLUÇÕES - {serie.upper()}", level=0)
        titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        subtitulo = document.add_heading(f"Tema: {tema}", level=1)
        subtitulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        document.add_paragraph("")
        document.add_paragraph("Resoluções detalhadas para conferência do professor.")
        document.add_paragraph("")
        
        # Processar questões
        introducao, questoes = parse_prova_completa(texto)
        
        for idx, questao in enumerate(questoes, start=1):
            # Título da questão
            p_gab = document.add_paragraph()
            p_gab.add_run(f"Questão {idx}: ").bold = True
            
            # Resolução
            if questao["resolucao"]:
                p_gab.add_run(questao["resolucao"])
            
            # Referência
            if questao["referencia"]:
                document.add_paragraph("")
                p_ref = document.add_paragraph()
                p_ref.add_run("Referência: ").bold = True
                p_ref.add_run(questao["referencia"])
            
            document.add_paragraph("")
        
        # Salvar
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        nome_arquivo = f"gabarito_{serie.replace(' ', '_')}_{tema.replace(' ', '_')}_{timestamp}.docx"
        document.save(nome_arquivo)
        
        return nome_arquivo
        
    except Exception as e:
        st.error(f"Erro ao criar gabarito: {str(e)}")
        return None

def criar_folha_respostas_otimizada(num_questoes, serie, tema):
    """Cria folha de respostas otimizada"""
    try:
        document = Document()
        
        # Cabeçalho
        titulo = document.add_heading("FOLHA DE RESPOSTAS", level=0)
        titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        subtitulo = document.add_paragraph(f"{serie} - {tema}")
        subtitulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        document.add_paragraph("")
        
        # Dados do aluno
        document.add_paragraph("Nome: _________________________________________________ Turma: _______ Data: ________")
        document.add_paragraph("")
        
        # Instruções
        document.add_paragraph("Instruções: Marque apenas uma alternativa por questão.")
        document.add_paragraph("")
        
        # Grid de respostas otimizado
        questoes_por_linha = min(5, num_questoes)
        linhas = (num_questoes + questoes_por_linha - 1) // questoes_por_linha
        
        for linha in range(linhas):
            p = document.add_paragraph()
            for col in range(questoes_por_linha):
                questao_num = linha * questoes_por_linha + col + 1
                if questao_num <= num_questoes:
                    p.add_run(f"{questao_num:2d}) ")
                    for letra in ['A', 'B', 'C', 'D']:
                        p.add_run(f"( {letra} ) ")
                    p.add_run("     ")
        
        # Salvar
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        nome_arquivo = f"folha_respostas_{serie.replace(' ', '_')}_{tema.replace(' ', '_')}_{timestamp}.docx"
        document.save(nome_arquivo)
        
        return nome_arquivo
        
    except Exception as e:
        st.error(f"Erro ao criar folha de respostas: {str(e)}")
        return None

# === PROMPT OTIMIZADO ===
def criar_prompt_otimizado(serie, tema, num_questoes, nivel_dificuldade):
    tempo_estimado = num_questoes * 6  # 6 min por questão
    
    return f"""
Crie uma prova de matemática completa para {serie} sobre o tema "{tema}" com {num_questoes} questões.

⚠️ FORMATAÇÃO MATEMÁTICA OBRIGATÓRIA:
- NÃO use códigos LaTeX como \\( \\) ou \\[ \\]
- Use símbolos Unicode diretos: ², ³, √, ÷, ×, ±, π, α, β
- Para frações use: (a)/(b) em vez de \\frac{{a}}{{b}}
- Para exponenciais use: x² ou x^2 (sem barras)
- Para coordenadas use: A(1, 3) sem delimitadores LaTeX
- Para raízes use: √25 ou raiz de 25

⚠️ REFERÊNCIAS OBRIGATÓRIAS:
- Cada questão DEVE ter uma referência bibliográfica específica
- Use apenas livros didáticos conhecidos
- Cite página específica quando possível
- Formate as referências em padrão ABNT

FORMATO OBRIGATÓRIO:

# Prova de Matemática - {serie}: {tema}

Instruções: Esta prova contém {num_questoes} questões sobre {tema}. Leia atentamente cada questão antes de responder. Marque apenas uma alternativa por questão.

Tempo estimado: {tempo_estimado} minutos

## Questão 1: [Tipo - Nível]
**Enunciado:** [Questão aqui - SEM códigos LaTeX, com medidas e dados específicos]
**Alternativas:**
a) [Opção A - formatação limpa]
b) [Opção B - formatação limpa] 
c) [Opção C - formatação limpa]
d) [Opção D - formatação limpa]
**Resolução:** [Explicação detalhada com cálculos limpos]
**Referência:** [Cite livro específico - formato ABNT completo]

[Repetir para todas as {num_questoes} questões]

## GABARITO E RESOLUÇÕES
[Resoluções detalhadas de cada questão]

## REFERÊNCIAS BIBLIOGRÁFICAS UTILIZADAS
[Liste todas as referências citadas, formatadas em ABNT]

CRITÉRIOS OBRIGATÓRIOS:
- Cada questão DEVE ter uma referência específica
- Formatação matemática LIMPA (sem LaTeX)
- Include medidas, ângulos, coordenadas específicas nas questões
- Varie as referências entre questões
- Use apenas símbolos Unicode ou texto simples
- Garanta que todas as expressões sejam legíveis
"""

def obter_resposta_completa_do_assistant(messages):
    """Extrai conteúdo completo das mensagens do assistant"""
    conteudo_completo = ""
    for message in messages.data:
        if message.role == "assistant":
            if hasattr(message, 'content') and message.content:
                for content_block in message.content:
                    if hasattr(content_block, 'text') and hasattr(content_block.text, 'value'):
                        conteudo_completo += content_block.text.value + "\n\n"
                    elif hasattr(content_block, 'text'):
                        conteudo_completo += str(content_block.text) + "\n\n"
    return conteudo_completo.strip()

# === INTERFACE PRINCIPAL ===
def main():
    st.title("📚 Gerador de Provas IA - Versão Corrigida")
    st.markdown("*Sistema inteligente com imagens contextuais e downloads sem reinicialização*")
    st.markdown("---")
    
    # Sidebar - Configurações
    st.sidebar.header("⚙️ Configurações da Prova")
    
    serie = st.sidebar.selectbox(
        "📖 Série:",
        ["6º Ano", "7º Ano", "8º Ano", "9º Ano", "1º Ano EM", "2º Ano EM", "3º Ano EM"]
    )
    
    # Usar a lista extensa de temas
    tema = st.sidebar.selectbox("💡 Tema:", TEMAS_COMPLETOS[serie])
    
    # Configurações da prova
    num_questoes = st.sidebar.slider("📊 Número de questões:", 3, 10, 5)
    nivel_dificuldade = st.sidebar.selectbox(
        "⚡ Nível de Dificuldade:",
        ["Fácil", "Médio", "Difícil", "Misto"]
    )
    
    tempo_estimado = num_questoes * 6
    
    # Opções visuais
    st.sidebar.markdown("### 🎨 Recursos Visuais")
    incluir_imagens = st.sidebar.checkbox("🖼️ Gerar imagens contextuais", value=True)
    incluir_folha_respostas = st.sidebar.checkbox("📝 Gerar folha de respostas", value=True)
    
    # Métricas
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("📊 Questões", num_questoes)
    with col2:
        st.metric("⏱️ Tempo", f"{tempo_estimado} min")
    with col3:
        st.metric("🎯 Nível", nivel_dificuldade)
    with col4:
        temas_disponiveis = len(TEMAS_COMPLETOS[serie])
        st.metric("💡 Temas", f"{temas_disponiveis} opções")
    
    # Botão principal de geração
    if st.button("🚀 Gerar Prova Completa", type="primary"):
        
        with st.spinner("🔄 Gerando prova inteligente..."):
            try:
                # Criar prompt otimizado
                prompt = criar_prompt_otimizado(serie, tema, num_questoes, nivel_dificuldade)
                
                # Iniciar conversa com assistente
                thread = openai.beta.threads.create()
                
                message = openai.beta.threads.messages.create(
                    thread_id=thread.id,
                    role="user", 
                    content=prompt
                )
                
                # Executar assistente
                run = openai.beta.threads.runs.create(
                    thread_id=thread.id,
                    assistant_id=ASSISTANT_ID
                )
                
                # Aguardar conclusão com progress bar
                progress_bar = st.progress(0)
                progress_text = st.empty()
                
                step = 0
                while run.status not in ["completed", "failed", "cancelled"]:
                    time.sleep(3)
                    run = openai.beta.threads.runs.retrieve(thread_id=thread.id, run_id=run.id)
                    step += 1
                    progress = min(step * 15, 95)
                    progress_bar.progress(progress)
                    progress_text.text(f"⏳ Processando: {run.status} - {progress}%")
                
                progress_bar.progress(100)
                progress_text.text("✅ Prova gerada com sucesso!")
                
                if run.status == "completed":
                    # Obter resposta do assistant
                    messages = openai.beta.threads.messages.list(thread_id=thread.id)
                    prova_gerada = obter_resposta_completa_do_assistant(messages)
                    
                    if prova_gerada and len(prova_gerada) > 100:
                        
                        # Limpar formatação matemática
                        prova_gerada = limpar_formatacao_latex(prova_gerada)
                        
                        # SALVAR NO SESSION STATE (CORREÇÃO DO BUG DE DOWNLOADS)
                        st.session_state.prova_gerada = prova_gerada
                        
                        st.success("✅ Prova gerada com sucesso!")
                        
                        # Preview
                        st.markdown("### 📖 Preview da Prova")
                        with st.expander("Ver conteúdo completo", expanded=False):
                            st.markdown(prova_gerada)
                        
                        # Gerar documentos
                        st.markdown("### 🔧 Gerando Documentos")
                        
                        with st.spinner("📄 Criando documentos..."):
                            # Gerar todos os documentos
                            nome_prova = criar_documento_prova_completo(
                                prova_gerada, serie, tema, incluir_imagens
                            )
                            
                            nome_gabarito = criar_documento_gabarito(
                                prova_gerada, serie, tema
                            )
                            
                            nome_folha = None
                            if incluir_folha_respostas:
                                nome_folha = criar_folha_respostas_otimizada(
                                    num_questoes, serie, tema
                                )
                        
                        # SALVAR DOCUMENTOS NO SESSION STATE
                        st.session_state.documentos_prontos = {
                            'prova': nome_prova,
                            'gabarito': nome_gabarito,
                            'folha': nome_folha
                        }
                        
                    else:
                        st.error("❌ Conteúdo gerado está vazio ou muito curto.")
                else:
                    st.error(f"❌ Erro na geração: {run.status}")
                    
            except Exception as e:
                st.error(f"❌ Erro: {str(e)}")
    
    # SEÇÃO DE DOWNLOADS (SEM REINICIALIZAÇÃO)
    if st.session_state.prova_gerada and st.session_state.documentos_prontos:
        st.markdown("### 📁 Downloads Disponíveis")
        
        col_down1, col_down2, col_down3 = st.columns(3)
        
        # Download da Prova
        with col_down1:
            if st.session_state.documentos_prontos['prova'] and os.path.exists(st.session_state.documentos_prontos['prova']):
                with open(st.session_state.documentos_prontos['prova'], "rb") as file_obj:
                    st.download_button(
                        label="📄 Baixar Prova",
                        data=file_obj.read(),
                        file_name=st.session_state.documentos_prontos['prova'],
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary"
                    )
        
        # Download do Gabarito
        with col_down2:
            if st.session_state.documentos_prontos['gabarito'] and os.path.exists(st.session_state.documentos_prontos['gabarito']):
                with open(st.session_state.documentos_prontos['gabarito'], "rb") as file_obj:
                    st.download_button(
                        label="📋 Baixar Gabarito",
                        data=file_obj.read(),
                        file_name=st.session_state.documentos_prontos['gabarito'],
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="secondary"
                    )
        
        # Download da Folha de Respostas
        with col_down3:
            if st.session_state.documentos_prontos['folha'] and os.path.exists(st.session_state.documentos_prontos['folha']):
                with open(st.session_state.documentos_prontos['folha'], "rb") as file_obj:
                    st.download_button(
                        label="📝 Baixar Folha Respostas",
                        data=file_obj.read(),
                        file_name=st.session_state.documentos_prontos['folha'],
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="secondary"
                    )
        
        # Informações finais
        st.success("✅ Todos os documentos estão prontos para download!")
        
        # Estatísticas da prova gerada
        introducao, questoes = parse_prova_completa(st.session_state.prova_gerada)
        num_questoes_geradas = len(questoes)
        imagens_geradas = len(st.session_state.imagens_geradas)
        
        st.info(f"""
        📊 **Resumo da Prova Gerada:**
        • {num_questoes_geradas} questões de {nivel_dificuldade.lower()} dificuldade
        • {imagens_geradas} imagens contextuais geradas
        • Tempo estimado: {tempo_estimado} minutos
        • Série: {serie} | Tema: {tema}
        """)
        
        # Botão para limpar e gerar nova prova
        if st.button("🔄 Gerar Nova Prova", type="secondary"):
            # Limpar session state
            st.session_state.prova_gerada = None
            st.session_state.documentos_prontos = None
            
            # Limpar imagens
            for img in st.session_state.imagens_geradas:
                if os.path.exists(img):
                    try:
                        os.remove(img)
                    except:
                        pass
            st.session_state.imagens_geradas = []
            
            st.rerun()
    
    # Footer
    st.markdown("---")
    st.markdown("""
    
        🤖 Gerador de Provas IA - Versão Corrigida
        ✨ Imagens contextuais • 📚 Referências ABNT • 🔧 Downloads sem reinicialização
    
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
